import docx

def generate_timetable():
    doc = docx.Document()
    doc.add_heading('Sample Timetable', 0)

    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    periods = ['P1', 'P2', 'P3', 'P4', 'P5', 'P6', 'P7']
    
    table = doc.add_table(rows=len(periods)+1, cols=len(days)+1)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    for i, day in enumerate(days):
        hdr_cells[i+1].text = day
        
    schedule = {
        'P1': ['Math - Dr. Smith', 'Physics - Prof. J', 'Chemistry - Dr. W', 'Break', 'CS - Dr. A', 'CS - Dr. A'],
        'P2': ['Break', 'Math - Dr. Smith', 'CS - Dr. A', 'Physics - Prof. J', 'Math - Dr. Smith', 'Break'],
        'P3': ['CS - Dr. A', 'Break', 'Math - Dr. Smith', 'Chemistry - Dr. W', 'Physics - Prof. J', 'Physics - Prof. J'],
        'P4': ['Physics - Prof. J', 'CS - Dr. A', 'Break', 'Math - Dr. Smith', 'Chemistry - Dr. W', 'Break'],
        'P5': ['Break', 'Break', 'Physics - Prof. J', 'CS - Dr. A', 'Math - Dr. Smith', 'CS - Dr. A'],
        'P6': ['Chemistry - Dr. W', 'Physics - Prof. J', 'CS - Dr. A', 'Break', 'Break', 'Math - Dr. Smith'],
        'P7': ['Math - Dr. Smith', 'Chemistry - Dr. W', 'Physics - Prof. J', 'CS - Dr. A', 'Physics - Prof. J', 'Break']
    }
    
    for row_idx, p in enumerate(periods):
        row_cells = table.rows[row_idx+1].cells
        row_cells[0].text = p
        
        for col_idx, text in enumerate(schedule[p]):
            row_cells[col_idx+1].text = text
            
    doc.save('c:\\projects\\fs-backend2\\sample_timetable.docx')
    print("Successfully created sample_timetable.docx")

if __name__ == '__main__':
    generate_timetable()
